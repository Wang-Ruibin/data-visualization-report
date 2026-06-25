from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(BASE_DIR, "..", "report", "images")
REPORT_DIR = os.path.join(BASE_DIR, "..", "report")
os.makedirs(REPORT_DIR, exist_ok=True)

doc = Document()

FONT_BODY = '宋体'
FONT_HEADING = '黑体'

style = doc.styles['Normal']
style.font.name = FONT_BODY
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.first_line_indent = Cm(0.74)

def set_font(run, name, size):
    run.font.name = name
    run.font.size = size
    run.element.rPr.rFonts.set(qn('w:eastAsia'), name)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.first_line_indent = Cm(0)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = FONT_HEADING
        run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_HEADING)
    return h

def add_para(text, bold=False, align=None, indent=True, size=Pt(12)):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74) if indent else Cm(0)
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_font(run, FONT_BODY, size)
    if bold:
        run.bold = True
    return p

def add_chart(image_path, caption, width=Inches(5.2)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    run.add_picture(image_path, width=width)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Cm(0)
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after = Pt(8)
    run = cap.add_run(caption)
    run.font.size = Pt(10)
    run.font.name = FONT_BODY
    run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
    run.italic = True

img = lambda name: os.path.join(IMG_DIR, name)

# ====== Title Page ======
for _ in range(3):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.first_line_indent = Cm(0)
run = title_p.add_run("《数据可视化技术》课程报告")
run.font.size = Pt(26)
run.bold = True
run.font.name = FONT_HEADING
run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_HEADING)

doc.add_paragraph()

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_p.paragraph_format.first_line_indent = Cm(0)
run = subtitle_p.add_run("中国互联网发展态势可视化分析\n（2010–2025）")
run.font.size = Pt(18)
run.font.name = FONT_HEADING
run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_HEADING)

for _ in range(5):
    doc.add_paragraph()

info_items = [
    ("学    院", "商学院"),
    ("专    业", "信息管理与信息系统"),
    ("学    号", "2308080119"),
    ("姓    名", "王睿彬"),
    ("课    程", "数据可视化技术"),
    ("学    期", "2025–2026学年第二学期"),
]

from docx.shared import Emu
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls

table = doc.add_table(rows=6, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for row_idx, (label, value) in enumerate(info_items):
    row = table.rows[row_idx]
    row.height = Cm(1.2)

    cell0 = row.cells[0]
    cell0.width = Cm(4)
    cell0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    cell0.paragraphs[0].paragraph_format.first_line_indent = Cm(0)
    cell0.paragraphs[0].paragraph_format.space_before = Pt(0)
    cell0.paragraphs[0].paragraph_format.space_after = Pt(0)
    run0 = cell0.paragraphs[0].add_run(f"{label}：")
    run0.font.size = Pt(15)
    run0.font.name = FONT_BODY
    run0.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)

    cell1 = row.cells[1]
    cell1.width = Cm(8)
    cell1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell1.paragraphs[0].paragraph_format.first_line_indent = Cm(0)
    cell1.paragraphs[0].paragraph_format.space_before = Pt(0)
    cell1.paragraphs[0].paragraph_format.space_after = Pt(0)
    run1 = cell1.paragraphs[0].add_run(value)
    run1.font.size = Pt(15)
    run1.font.name = FONT_BODY
    run1.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)

    tc = cell1._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    tcBorders.append(bottom)
    tcPr.append(tcBorders)

for row in table.rows:
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.find(qn('w:tcBorders'))
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
        for side in ['top', 'left', 'right']:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'none')
            border.set(qn('w:sz'), '0')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
            tcBorders.append(border)
        tcPr.append(tcBorders)

doc.add_page_break()

# ====== 正文 ======

add_heading_styled("一、引言", level=1)

add_para(
    "从2010年到现在，中国互联网走过了波澜壮阔的十五年。"
    "记得2010年那会儿，3G网络刚刚普及不久，智能手机还是个新鲜玩意儿，"
    "上网主要还得靠电脑。再看看现在，人手一部手机就能搞定几乎所有事情——"
    "聊天、购物、看视频、点外卖、交水电费，互联网已经完全融入了日常生活的方方面面。"
    "这十几年里，中国互联网到底经历了怎样的变化？网民规模增长了多少？"
    "大家都是用什么设备上网的？上网的人都是什么年龄和学历？"
    "城市和农村之间的差距又有多大？这些问题如果能让数据来回答，应该会很有意思。"
)

add_para(
    "中国互联网络信息中心（CNNIC）每半年发布一次《中国互联网络发展状况统计报告》，"
    "里面记录了各种各样的互联网发展指标。本报告就是基于这些数据，"
    "把2010年到2025年间的关键指标整理出来，用图表的形式做可视化分析。"
    "具体的维度包括网民规模、互联网普及率、手机网民占比、网民年龄和学历结构、"
    "各类应用的用户规模、城乡普及率对比，以及IP地址资源的变化。"
)

add_para(
    "可视化工具方面，我用了PyECharts和Matplotlib两个库。"
    "PyECharts生成的是交互式网页图表，用浏览器打开以后可以悬浮看数据、缩放、筛选，"
    "交互性比较强。报告中插入的静态图表则用Matplotlib绘制，方便打印和阅读。"
    "两种工具各有所长，搭配着用效果更好。"
)

add_heading_styled("二、数据来源与处理", level=1)

add_para(
    "本报告的数据全部来自CNNIC第27次到第55次《中国互联网络发展状况统计报告》，"
    "时间跨度从2010年到2025年。具体来说，包括七个数据集："
    "网民规模与互联网普及率（2010至2025年，逐年数据）、"
    "手机网民规模及占比（2010至2025年，逐年数据）、"
    "网民年龄结构（2013、2016、2019、2022、2024年）、"
    "网民学历结构（2013、2016、2019、2022、2024年）、"
    "各类互联网应用用户规模（2024年）、"
    "城镇与农村互联网普及率对比（2013、2016、2019、2022、2024年），"
    "以及IPv4/IPv6地址资源数量（2017至2024年，逐年数据）。"
)

add_para(
    "数据处理过程并不复杂：从CNNIC报告原文中摘录数字，整理成CSV格式，"
    "再用Python的Pandas库读取。对于那些不是每年都有的指标（比如年龄结构和学历结构），"
    "我直接用了现有年份的数据，没有做任何插值或推算，确保每个数字都是报告中的原始值。"
)

add_heading_styled("三、可视化分析与发现", level=1)

add_heading_styled("3.1 网民规模和普及率——从4.5亿到11亿的飞跃", level=2)

add_chart(img("chart1_users_penetration.png"),
          "图1 中国网民规模与互联网普及率变化（2010-2025）")

add_para(
    "图1展示的是中国网民规模和互联网普及率从2010年到2025年的变化趋势。"
    "从图中可以直观地看到，网民规模从2010年的4.57亿增长到了2025年的11.25亿，"
    "十五年里翻了一倍多。普及率也从34.3%提高到了79.8%，"
    "也就是说到现在，将近八成中国人都在上网了。"
)

add_para(
    "仔细看增长曲线的话，前十年（2010到2020年）增长非常迅猛，"
    "基本上每年新增四五千万网民。这段时间正好赶上了3G普及、4G爆发和智能手机价格迅速下降，"
    "上网的门槛一下子降低了很多。2020年之后增速明显放缓，一年也就新增两三千万人，"
    "这也符合客观规律——能上网的人大部分都已经上了，市场逐渐饱和，"
    "接下来更多是存量竞争而不是增量扩张了。"
)

add_heading_styled("3.2 手机网民——从配角到绝对主角", level=2)

add_chart(img("chart2_mobile.png"),
          "图2 中国手机网民规模及占比变化（2010-2025）")

add_para(
    "图2反映的是手机上网的情况。2010年的时候，虽然已经有了手机网民这个概念，"
    "但占比只有66.2%，还有三分之一的人是用电脑或其他设备上网的。"
    "到了2025年，手机网民占比已经达到99.9%，基本上可以说是全民手机上网了。"
)

add_para(
    "一个关键的转折点在2015年，那一年手机网民占比首次突破90%。"
    "从那以后，桌面端的存在感就越来越弱。现在做一个互联网产品如果没有手机版，"
    "基本等于放弃整个市场。这种移动化的趋势很大程度上塑造了中国互联网的生态——"
    "中国的移动支付、短视频、外卖等服务之所以能走在世界前列，"
    "和全民手机上网的格局是分不开的。"
)

add_heading_styled("3.3 年龄结构——互联网不再是年轻人的专属", level=2)

add_chart(img("chart3_age.png"),
          "图3 中国网民年龄结构演变（2013-2024）")

add_para(
    "图3用堆叠柱状图展示了不同年龄段网民占比的变化。这个图透露出的信息很丰富。"
)

add_para(
    "最大的变化来自中老年群体。2013年时，50岁以上的网民加在一起只占7.6%，"
    "几乎可以忽略。但到了2024年，这个比例涨到了22.8%，"
    "特别是60岁以上的老年网民，从2.5%增长到了7.6%。"
    "想想身边的亲戚长辈，是不是这几年也开始刷短视频、用微信发红包了？"
    "数据正好反映了这个趋势。"
)

add_para(
    "年轻网民的比例则是在下降的。20到29岁年龄段从31.2%降到了18.7%，"
    "降了将近13个百分点。不过这不代表年轻人不上网了，"
    "更可能的原因是中老年群体大量涌入，把年轻群体的比例稀释了。"
    "实际上年轻网民的上网时长和活跃度仍然是最高的。"
    "30到39岁群体占比保持稳定，始终在20%到24%之间，"
    "是互联网用户的中坚力量。40到49岁群体从12.5%增长到21.5%，"
    "也说明互联网正在向更年长的群体渗透。"
)

add_heading_styled("3.4 学历结构——分布越来越均匀", level=2)

add_chart(img("chart4_education.png"),
          "图4 2024年中国网民学历结构分布")

add_para(
    "图4是2024年网民学历结构的环形图。初中学历的网民占比最高，达到32.8%，"
    "高中或中专占26.2%，大专或本科占27.6%，三者相差不大。"
    "小学及以下占8.8%，硕士及以上占4.6%。整体分布相对均匀。"
)

add_para(
    "从趋势上看，高学历网民的比例在稳步提升。"
    "大专或本科及以上学历的占比从2013年的20.9%增长到了2024年的32.2%，"
    "说明互联网正在吸引越来越多受过高等教育的人。"
    "不过低学历群体的绝对规模仍然很大，"
    "这意味着互联网产品需要在设计上兼顾不同教育背景的用户——"
    "并不是所有人都能轻松理解复杂的操作界面。"
)

add_heading_styled("3.5 互联网应用——即时通信和视频领跑", level=2)

add_chart(img("chart5_apps.png"),
          "图5 2024年中国主要互联网应用用户规模排名")

add_para(
    "图5是2024年各类互联网应用的用户规模排名。即时通信以10.7亿用户排名第一，"
    "微信和QQ几乎覆盖了所有网民。网络视频（含短视频）以10.5亿紧随其后，"
    "抖音、快手等短视频平台的贡献功不可没。网络支付以9.8亿排第三，"
    "差不多是网民就会用移动支付了。搜索引擎（8.9亿）和网络购物（8.7亿）"
    "分列第四和第五位，都是名副其实的国民级应用。"
)

add_para(
    "比较值得关注的是在线教育（4.2亿）和在线办公（5.3亿）。"
    "这两类应用在2020年疫情期间经历了一波爆发式增长，"
    "疫情之后虽然有所回落，但用户规模仍然显著高于疫情前的水平。"
    "这说明有些使用习惯一旦养成，就很难再回到原来的方式了。"
    "网络文学的5.5亿用户也表明数字内容消费已经成为网民日常生活的重要组成部分。"
)

add_heading_styled("3.6 城乡差距——进步很大，差距仍在", level=2)

add_chart(img("chart6_urban_rural.png"),
          "图6 中国城乡互联网普及率对比（2013-2024）")

add_para(
    "图6对比了城镇和农村的互联网普及率。城镇从2013年的62.0%增长到了2024年的84.5%，"
    "农村从28.1%增长到了65.3%。虽然农村的绝对水平还赶不上城镇，"
    "但增长速度明显更快，十年间从28%涨到65%，增幅超过130%。"
)

add_para(
    "城乡差距从2013年的34个百分点缩小到了2024年的19个百分点，"
    "说明数字鸿沟正在逐渐收窄，但将近20个百分点的差距依然不小。"
    "这背后涉及网络基础设施、经济水平、教育程度等多方面因素。"
    "这几年国家推进的宽带中国战略和电信普遍服务工作确实起到了积极作用，"
    "但要让城乡互联网普及率完全拉平，还需要长期努力。"
)

add_heading_styled("3.7 IP地址资源——IPv6正在加速追赶", level=2)

add_chart(img("chart7_ip.png"),
          "图7 中国IPv4/IPv6地址资源变化（2017-2024）")

add_para(
    "最后一个图表涉及的是IP地址资源。IPv4地址这几年几乎没有增长，"
    "从3.39亿到3.62亿，涨幅很小。原因很简单——全球IPv4地址早就分配完了，"
    "用一点少一点。IPv6的情况就完全不同了。"
    "从2017年的0.86亿增长到了2024年的2.63亿，翻了将近三倍。"
    "这和国家政策的推动密不可分。2017年国家专门发文部署IPv6规模部署，"
    "运营商和互联网企业积极响应，现在中国的IPv6活跃用户数在全球名列前茅。"
    "未来IPv6全面取代IPv4应该只是时间问题。"
)

add_heading_styled("四、总结", level=1)

add_para(
    "回顾2010年到2025年中国互联网这十五年的发展历程，可以看到几个很清晰的趋势。"
    "首先，网民规模从4.57亿增长到了11.25亿，互联网普及率接近八成，"
    "但增量红利基本见顶，未来更多是存量竞争。其次，手机已经成为绝对主流的上网终端，"
    "手机网民占比超过99.9%，移动优先是所有互联网服务的基本前提。"
    "第三，用户结构越来越多元，老人和小孩都在上网，"
    "学历分布也越来越均匀，互联网服务需要兼顾不同群体的使用习惯。"
    "第四，城乡差距在持续缩小，农村普及率十年翻了一番多，"
    "但跟城镇还有大约20个百分点的差距，数字鸿沟问题仍然值得关注。"
    "第五，基础设施在不断升级，IPv6的快速部署为互联网的长期发展打下了良好基础。"
    "最后，互联网应用生态非常丰富，从聊天、购物到看视频、办公，"
    "互联网已经渗透到了社会生活的每一个角落。"
)

add_para(
    "做这个报告本身也是一个学习和思考的过程。以前看CNNIC的报告觉得就是一堆数字，"
    "但真正动手把数据整理成图表之后，很多藏在数字背后的趋势就一目了然了。"
    "数据可视化确实是一个很好的工具，它让数据自己说话，"
    "让复杂的信息变得直观和易于理解。"
)

add_heading_styled("参考文献", level=1)

refs = [
    "[1] 中国互联网络信息中心. 第27次中国互联网络发展状况统计报告[R]. 北京, 2010.",
    "[2] 中国互联网络信息中心. 第37次中国互联网络发展状况统计报告[R]. 北京, 2015.",
    "[3] 中国互联网络信息中心. 第47次中国互联网络发展状况统计报告[R]. 北京, 2020.",
    "[4] 中国互联网络信息中心. 第51次中国互联网络发展状况统计报告[R]. 北京, 2022.",
    "[5] 中国互联网络信息中心. 第53次中国互联网络发展状况统计报告[R]. 北京, 2023.",
    "[6] 中国互联网络信息中心. 第55次中国互联网络发展状况统计报告[R]. 北京, 2024.",
    "[7] 中共中央办公厅, 国务院办公厅. 推进互联网协议第六版（IPv6）规模部署行动计划[Z]. 2017.",
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(ref)
    set_font(run, FONT_BODY, Pt(11))

out_path = os.path.join(REPORT_DIR, "2308080119+王睿彬+数据源码.docx")
doc.save(out_path)
print(f"Report saved to: {out_path}")
